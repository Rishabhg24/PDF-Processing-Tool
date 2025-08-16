import os
import sys
import subprocess
import fitz
import pdfplumber
from PyPDF2 import PdfMerger
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from tkinter import filedialog, messagebox
import ttkbootstrap as tb
import tkinter as tk
from ttkbootstrap.constants import *
from pdf2docx import Converter
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl import load_workbook
from docx.api import Document as DocxDocument
import time
from PIL import Image
import io
import uuid
import hashlib
from checksum_validator import validate_checksum, generate_data_file, calculate_checksum, save_checksum
import ctypes

try:
    ctypes.windll.shcore.SetProcessDpiAwareness(1)
except Exception:
    pass

# Set taskbar icon for Windows
try:
    if sys.platform == "win32":
        # Set the application ID for the taskbar
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID("FutureWave.PDFProcessingTool.1")
except Exception:
    pass

def validate_application():
    """Validate the application's integrity"""
    status = validate_checksum()
    if status == 'corrupt':
        messagebox.showerror("Integrity Error", "Application data has been modified or corrupted. Please reinstall.")
        sys.exit(1)
    elif status == 'expired':
        sys.exit(1)

# Add validation check at startup
validate_application()

def get_exe_path():
    if getattr(sys, 'frozen', False):  # EXE mode
        return sys._MEIPASS
    return os.path.dirname(os.path.abspath(__file__))

def get_mac_address():
    mac = uuid.getnode()
    return ':'.join(("%012X" % mac)[i:i+2] for i in range(0, 12, 2))

#print("MAC Address:", get_mac_address())

def generate_license_key(mac_address):
    secret = "FUTUREWAVE_SECRET"
    raw = mac_address + secret
    return hashlib.sha256(raw.encode()).hexdigest().upper()

mac = "E0:2B:E9:49:75:31"
#print(generate_license_key(mac))

def validate_license():
    expected_key = generate_license_key(get_mac_address())
    key_file = "license.key"

    if not os.path.exists(key_file):
        messagebox.showerror("Access Denied", "License file not found.")
        sys.exit()

    with open(key_file, "r") as f:
        license_key = f.read().strip()

    if license_key != expected_key:
        messagebox.showerror("Access Denied", "Invalid license key for this machine.")
        sys.exit() 


def get_mac_address():
    mac = uuid.getnode()
    return ':'.join(("%012X" % mac)[i:i+2] for i in range(0, 12, 2))

def update_custom_name_state():
    files = pdf_files_var.get()
    if len(files) > 1:
        custom_name_entry.config(state="disabled")
    else:
        custom_name_entry.config(state="normal")


def add_more_pdfs():
    new_files = filedialog.askopenfilenames(filetypes=[("PDF Files", "*.pdf")])
    if new_files:
        current_files = list(pdf_files_var.get())
        combined_files = current_files + list(new_files)
        # Remove duplicates while preserving order
        seen = set()
        unique_files = [f for f in combined_files if not (f in seen or seen.add(f))]
        pdf_files_var.set(unique_files)
        refresh_file_listbox()
        update_custom_name_state()



def move_selected_up():
    selected = file_listbox.curselection()
    if not selected or selected[0] == 0:
        return
    idx = selected[0]
    files = list(pdf_files_var.get())
    files[idx - 1], files[idx] = files[idx], files[idx - 1]
    pdf_files_var.set(files)
    refresh_file_listbox()
    file_listbox.select_set(idx - 1)

def move_selected_down():
    selected = file_listbox.curselection()
    if not selected or selected[0] == file_listbox.size() - 1:
        return
    idx = selected[0]
    files = list(pdf_files_var.get())
    files[idx + 1], files[idx] = files[idx], files[idx + 1]
    pdf_files_var.set(files)
    refresh_file_listbox()
    file_listbox.select_set(idx + 1)

def refresh_file_listbox():
    file_listbox.delete(0, tk.END)
    for f in pdf_files_var.get():
        file_listbox.insert(tk.END, os.path.basename(f))

def clear_all_files():
    pdf_files_var.set([])
    file_listbox.delete(0, tk.END)
    update_custom_name_state()


# compression function
def Compressed_pdfs(progress_callback=None):
    selected_option = radio_var.get()
    files = pdf_files_var.get()
    output_folder = output_folder_var.get()

    if not files:
        messagebox.showerror("No PDFs", "Please select PDF files to compress.")
        return

    if not output_folder:
        output_folder = os.path.dirname(files[0])

    compressed_paths = []
    non_compressible_files = []

    # First, count total pages across all PDFs for better granularity
    total_pages = 0
    for f in files:
        try:
            total_pages += fitz.open(f).page_count
        except:
            continue

    pages_processed = 0

    for file_path in files:
        try:
            # Check if file is already compressed or too small
            file_size = os.path.getsize(file_path)
            if file_size < 1024 * 50:  # If file is less than 50KB
                non_compressible_files.append((file_path, "File is too small to compress"))
                continue

            src_doc = fitz.open(file_path)
            
            # Check if PDF is already compressed
            is_compressed = True
            for page in src_doc:
                if page.get_images():  # If page has images
                    is_compressed = False
                    break
            
            if is_compressed:
                non_compressible_files.append((file_path, "PDF appears to be already compressed"))
                src_doc.close()
                continue

            new_doc = fitz.open()
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            suffix = "Extreme" if selected_option == "Extreme Compress (Low Quality,High compression)" else "low"
            output_path = os.path.join(output_folder, f"{base_name}_{suffix}_compressed.pdf")

            # Set compression parameters
            if selected_option == "Extreme Compress (Low Quality,High compression)":
                jpeg_quality = 15
                dpi_limit = 100
                min_size = 100
            else:
                jpeg_quality = 60
                dpi_limit = 200
                min_size = 50

            for page in src_doc:
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                if max(pix.width, pix.height) > dpi_limit * 10:
                    scale = (dpi_limit * 10) / max(pix.width, pix.height)
                    new_size = (int(pix.width * scale), int(pix.height * scale))
                    img = img.resize(new_size, Image.LANCZOS)

                img_buffer = io.BytesIO()
                img.save(img_buffer, format="JPEG", quality=jpeg_quality)
                img_bytes = img_buffer.getvalue()

                rect = fitz.Rect(0, 0, img.width, img.height)
                new_page = new_doc.new_page(width=img.width, height=img.height)
                new_page.insert_image(rect, stream=img_bytes)

                pages_processed += 1
                if progress_callback and total_pages:
                    percent = int((pages_processed / total_pages) * 100)
                    progress_callback(percent)
                    root.update_idletasks()

            new_doc.save(output_path)
            new_doc.close()
            src_doc.close()

            if os.path.getsize(output_path) < os.path.getsize(file_path):
                compressed_paths.append(output_path)
            else:
                os.remove(output_path)
                non_compressible_files.append((file_path, "Compression did not reduce file size"))

        except Exception as e:
            non_compressible_files.append((file_path, f"Error: {str(e)}"))
            if progress_callback:
                progress_callback(0)
            continue

    # Show messages for non-compressible files
    if non_compressible_files:
        non_compressible_msg = "The following files could not be compressed:\n\n"
        for file_path, reason in non_compressible_files:
            non_compressible_msg += f"{os.path.basename(file_path)}: {reason}\n"
        messagebox.showwarning("Compression Warning", non_compressible_msg)

    if compressed_paths:
        messagebox.showinfo("Compression Complete", "Compressed files saved:\n" + "\n".join(compressed_paths))
        open_output_folder(output_folder)
    elif not non_compressible_files:  # If no files were compressed and no warnings were shown
        messagebox.showwarning("No Files Compressed", "No files were compressed successfully.")

    if progress_callback:
        progress_callback(0)  # Reset progress at the end

# --- Merge Functions ---
def merge_pdfs(pdf_paths, output_path):
    try:
        merger = PdfMerger()
        for pdf in pdf_paths:
            merger.append(pdf)
        merger.write(output_path)
        merger.close()
        return True, output_path
    except Exception as e:
        return False, str(e)

def merge_selected_pdfs():
    files = pdf_files_var.get()
    output_folder = output_folder_var.get()

    if not files:
        messagebox.showerror("No PDFs", "Please select PDF files to merge.")
        return

    if not output_folder:
        output_folder = os.path.dirname(files[0])

    base_name = custom_name_var.get().strip() or "merged"
    output_path = os.path.join(output_folder, f"{base_name}_merged.pdf")

    success, result = merge_pdfs(files, output_path)

    if success:
        messagebox.showinfo("Merge Complete", f"Merged PDF saved to:\n{result}")
        open_output_folder(output_folder)
    else:
        messagebox.showerror("Merge Failed", f"Error merging PDFs:\n{result}")

def pdf_to_word_full(pdf_path, output_folder, custom_name, progress_callback):
    if not os.path.exists(pdf_path):
        return False, "File not found"

    base_name = custom_name if custom_name else os.path.splitext(os.path.basename(pdf_path))[0]
    docx_path = os.path.join(output_folder, f"{base_name}.docx")

    try:
        progress_callback(10)
        root.update_idletasks() 
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        progress_callback(100)
        return True, docx_path
    except Exception as e:
        return False, str(e)

def word_to_excel_layout(docx_path, excel_path, progress_callback):
    doc = DocxDocument(docx_path)
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "PDF Content"

    row_idx = 1
    for block in doc.element.body:
        if block.tag.endswith('tbl'):
            table = next((tbl for tbl in doc.tables if tbl._tbl == block), None)
            if table:
                for row in table.rows:
                    col_idx = 1
                    for cell in row.cells:
                        ws.cell(row=row_idx, column=col_idx).value = cell.text.strip()
                        col_idx += 1
                    row_idx += 1
                row_idx += 1  # Space after table
        elif block.tag.endswith('p'):
            para = next((p for p in doc.paragraphs if p._p == block), None)
            if para and para.text.strip():
                ws.cell(row=row_idx, column=1).value = para.text.strip()
                row_idx += 1

    progress_callback(100)
    wb.save(excel_path)
    return excel_path

def pdf_to_excel_layout(pdf_path, output_folder, custom_name, progress_callback):
    success, docx_path_or_error = pdf_to_word_full(pdf_path, output_folder, custom_name, progress_callback)
    if not success:
        return False, docx_path_or_error

    base_name = custom_name if custom_name else os.path.splitext(os.path.basename(pdf_path))[0]
    excel_path = os.path.join(output_folder, f"{base_name}.xlsx")
    try:
        word_to_excel_layout(docx_path_or_error, excel_path, progress_callback)
        # Clean up the intermediate Word file
        try:
            os.remove(docx_path_or_error)
        except:
            pass  # Ignore if file deletion fails
        return True, excel_path
    except Exception as e:
        # Clean up the intermediate Word file even if conversion fails
        try:
            os.remove(docx_path_or_error)
        except:
            pass  # Ignore if file deletion fails
        return False, str(e)

#check_license_key()  # Add this before root = tb.Window(...)

# --- Helpers ---
def open_output_folder(path):
    if sys.platform == "win32":
        os.startfile(path)
    elif sys.platform == "darwin":
        subprocess.run(["open", path])
    else:
        subprocess.run(["xdg-open", path])

def is_file_open(filepath):
    try:
        # Try opening for reading and writing without truncating
        with open(filepath, 'a'):
            return False  # File is NOT open elsewhere
    except OSError:
        return True  # File IS open

def update_progress(target_value):
    if target_value == 0:
        progress_var.set(target_value)
    else:
        current = progress_var.get()
        step = 1 if target_value > current else -1
        for val in range(current, target_value + step, step):
            progress_var.set(val)
            root.update_idletasks()
            time.sleep(0.005)  # You can tweak this for speed (e.g., 0.005 is faster)

def select_pdfs():
    files = filedialog.askopenfilenames(filetypes=[("PDF Files", "*.pdf")])
    #if files:
        #pdf_files_var.set(files)
        #file_listbox.delete(0, tb.END)
        #for f in files:
            #file_listbox.insert(tb.END, os.path.basename(f))
    if files:
       pdf_files_var.set(files)
       refresh_file_listbox()
       update_custom_name_state()

def select_output_folder():
    folder = filedialog.askdirectory()
    if folder:
        output_folder_var.set(folder)

def convert_files(to_format):
    files = pdf_files_var.get()
    output_folder = output_folder_var.get()
    custom_name = custom_name_var.get().strip()

    if not files:
        messagebox.showerror("No PDFs", "Please select or drag-and-drop PDF files.")
        return

    if not output_folder:
        output_folder = os.path.dirname(files[0])

    update_progress(0)
    success_paths = []

    for file_path in files:
        name = custom_name if len(files) == 1 else ""
        if to_format == "excel":
            success, result = pdf_to_excel_layout(file_path, output_folder, name, update_progress)
        else:
            success, result = pdf_to_word_full(file_path, output_folder, name, update_progress)

        if success:
            success_paths.append(result)

    if success_paths:
        msg = "\n".join(success_paths)
        messagebox.showinfo("Done", f"Files saved:\n{msg}")
        open_output_folder(output_folder)
    else:
        messagebox.showwarning("Nothing converted", f"Unable to write output file(s), please check if it is already opened.")

    update_progress(0)  # 🔄 Reset the progress bar to 0

validate_license()

# --- UI Setup ---
root = tb.Window(themename="darkly")  # Changed from minty to darkly for better contrast

# Set icon for both window and taskbar
try:
    # Get the correct path to the icon file
    if getattr(sys, 'frozen', False):  # EXE mode
        icon_path = os.path.join(sys._MEIPASS, "icon.ico")
    else:  # Script mode
        icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "icon.ico")
    
    # Check if icon file exists
    if os.path.exists(icon_path):
        root.iconbitmap(icon_path)
        # Also set the taskbar icon explicitly
        if sys.platform == "win32":
            root.iconbitmap(default=icon_path)
    else:
        print(f"Icon file not found at: {icon_path}")
except Exception as e:
    print(f"Error setting icon: {e}")

root.title("PDF Processing Tool")
root.geometry("1100x850")
root.resizable(False, False)

# Create a style object for custom styling
style = tb.Style()
style.configure("Custom.TLabel", font=("Segoe UI", 9, "bold"), foreground="#FFFFFF")
style.configure("Custom.TButton", font=("Segoe UI", 8))
style.configure("Custom.TEntry", font=("Segoe UI", 8))

pdf_files_var = tb.Variable()
output_folder_var = tb.StringVar()
custom_name_var = tb.StringVar()
progress_var = tb.IntVar()

# Main container frame with padding
main_frame = tb.Frame(root, padding=20)
main_frame.pack(fill=BOTH, expand=True)

# File selection section
tb.Label(main_frame, text="Selected PDF files:", style="Custom.TLabel").pack(pady=(0, 5))
file_frame = tb.Frame(main_frame)
file_frame.pack(fill=X)

file_listbox = tk.Listbox(file_frame, width=70, height=6, bg="#2B2B2B", fg="#FFFFFF", 
                         selectbackground="#007BFF", selectforeground="#FFFFFF",
                         font=("Segoe UI", 8))
file_listbox.pack(side=LEFT, fill=X, expand=True)

arrow_frame = tb.Frame(file_frame)
arrow_frame.pack(side=LEFT, padx=(5, 0), anchor="n")

tk.Button(arrow_frame, text="🔼", font=("Arial", 18), bg="#2B2B2B", fg="#FFFFFF",
          activebackground="#007BFF", command=move_selected_up).pack(pady=(0, 5))
tk.Button(arrow_frame, text="🔽", font=("Arial", 18), bg="#2B2B2B", fg="#FFFFFF",
          activebackground="#007BFF", command=move_selected_down).pack()

# Browse and Clear buttons frame
browse_clear_frame = tb.Frame(main_frame)
browse_clear_frame.pack(pady=10)

# Create a container frame for the buttons with proper spacing
button_container = tb.Frame(browse_clear_frame)
button_container.pack(expand=True)

tb.Button(button_container, text="Browse PDFs", bootstyle="primary", 
          style="Custom.TButton", width=15, command=select_pdfs).pack(side=LEFT, padx=5)
tb.Button(button_container, text="Add More", bootstyle="info", 
          style="Custom.TButton", width=15, command=add_more_pdfs).pack(side=LEFT, padx=5)
tb.Button(button_container, text="Clear All", bootstyle="danger", 
          style="Custom.TButton", width=15, command=clear_all_files).pack(side=LEFT, padx=5)

# Output folder section
tb.Label(main_frame, text="Select Output Folder:", style="Custom.TLabel").pack(pady=(15, 5))
folder_frame = tb.Frame(main_frame)
folder_frame.pack(fill=X)
tb.Entry(folder_frame, textvariable=output_folder_var, width=50, 
         style="Custom.TEntry").pack(side=LEFT, fill=X, expand=True)
tb.Button(folder_frame, text="Browse", bootstyle="info", 
          style="Custom.TButton", command=select_output_folder).pack(side=LEFT, padx=5)

# Custom name section
tb.Label(main_frame, text="Custom File Name (optional):", style="Custom.TLabel").pack(pady=(15, 5))
custom_name_entry = tb.Entry(main_frame, textvariable=custom_name_var, width=40, 
                            style="Custom.TEntry")
custom_name_entry.pack(pady=5)

# Action buttons frame
button_frame = tb.Frame(main_frame)
button_frame.pack(pady=(20, 20))
tb.Button(button_frame, text="Convert to Excel", bootstyle="primary", width=25,
          style="Custom.TButton", command=lambda: convert_files("excel")).pack(side=LEFT, padx=10)
tb.Button(button_frame, text="Convert to Word", bootstyle="success", width=25,
          style="Custom.TButton", command=lambda: convert_files("word")).pack(side=LEFT, padx=10)
tb.Button(button_frame, text="Merge PDFs", bootstyle="warning", width=25,
          style="Custom.TButton", command=merge_selected_pdfs).pack(side=LEFT, padx=10)

# Compression options
compression_frame = tb.Frame(main_frame)
compression_frame.pack(pady=(10, 10))
radio_var = tk.StringVar(value="Less Compression (High Quality,less compression)")
options = ["Less Compression (High Quality,less compression)", 
           "Extreme Compress (Low Quality,High compression)"]
for option in options:
    tk.Radiobutton(compression_frame, text=option, variable=radio_var, value=option,
                   bg="#2B2B2B", fg="#FFFFFF", selectcolor="#2B2B2B",
                   activebackground="#2B2B2B", activeforeground="#FFFFFF").pack(anchor="w")

tb.Button(compression_frame, text="Compress file", bootstyle="info",
          style="Custom.TButton", command=lambda: Compressed_pdfs(progress_callback=update_progress)).pack(pady=10)

# Progress bar with enhanced visibility
progressbar = tb.Progressbar(main_frame, variable=progress_var, length=400, 
                            bootstyle="info-striped", style="Custom.TProgressbar")
progressbar.pack(pady=15)

# Support label with better contrast
support_label = tb.Label(main_frame, text="Support: Rishabh Private Limited |Rishabhg342@gmail.com", 
                        font=("Segoe UI", 8), foreground="#CCCCCC")
support_label.pack(side=BOTTOM, pady=5)

root.mainloop()
